"""Word 文档加载器。

支持 .docx 和 .doc 格式。
.docx 用 python-docx，.doc 依次尝试 antiword / catdoc / win32com。
"""

import os
import subprocess
import tempfile
from typing import Dict, List

from rag_server.registry import register
from rag_server.utils.timer import timeit
from .base import BaseLoader, LoaderOutput


@register("loader", "docx")
class WordLoader(BaseLoader):
    """Word 文档加载器，支持 .docx 和 .doc 格式。"""

    extensions = [".docx", ".doc"]

    @timeit
    async def load(self, file_path: str) -> LoaderOutput:
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".doc":
            text = await self._read_doc(file_path)
        else:
            from docx import Document
            doc = Document(file_path)
            text = self._extract_from_docx(doc)

        return LoaderOutput(
            md_text=text,
            assets=[],
            metadata={"source": file_path, "format": ext[1:]},
        )

    def _extract_from_docx(self, doc) -> str:
        """从 python-docx Document 提取 MD 格式文本。"""
        md_parts: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name.lower() if para.style else ""
                if "heading" in style_name or "title" in style_name:
                    level = style_name.replace("heading", "").strip()
                    if level.isdigit():
                        md_parts.append(f"{'#' * int(level)} {text}")
                    else:
                        md_parts.append(f"## {text}")
                else:
                    md_parts.append(text)
        return "\n\n".join(md_parts)

    async def _read_doc(self, file_path: str) -> str:
        """读取 .doc 格式，先后尝试 antiword / catdoc / win32com。"""
        # 1. antiword（最轻量）
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, timeout=30,
            )
            out = result.stdout.decode("utf-8", errors="replace").strip() if result.stdout else ""
            if result.returncode == 0 and out:
                return out
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # 2. catdoc
        try:
            result = subprocess.run(
                ["catdoc", file_path],
                capture_output=True, timeout=30,
            )
            out = result.stdout.decode("utf-8", errors="replace").strip() if result.stdout else ""
            if result.returncode == 0 and out:
                return out
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # 3. win32com（仅 Windows + MS Word）
        if os.name == "nt":
            tmp_path = None
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    tmp_path = tmp.name
                doc = word.Documents.Open(file_path)
                doc.SaveAs2(tmp_path, FileFormat=16)
                doc.Close()
                word.Quit()

                from docx import Document
                text = self._extract_from_docx(Document(tmp_path))
                return text
            except Exception:
                pass
            finally:
                if tmp_path:
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass

        raise ValueError(
            f"无法读取 .doc 文件: {file_path}。"
            f"请安装 antiword (https://www.winfield.demon.nl/) "
            f"或将文件另存为 .docx 格式。"
        )